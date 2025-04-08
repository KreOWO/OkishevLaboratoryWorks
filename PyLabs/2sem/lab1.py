import random
import os
from docx import Document
from openpyxl import Workbook, load_workbook
from fpdf import FPDF
import csv
import sqlite3


def create_data(n):
    student_names_m = ['Кирилл', 'Андрей', 'Пётр', 'Даниил', 'Григорий', 'Филипп', 'Ян', 'Ибрагим', 'Алексей',
                       'Евгений', 'Антон']
    student_names_w = ['Ольга', 'Ника', 'Валерия', 'Надежда', 'Анна', 'Василиса', 'Яна', 'Любовь', 'Ангелина',
                       'Эвелина', 'Елизавета']
    student_names = [student_names_m, student_names_w]
    student_snames = ['Смирнов', 'Иванов', 'Сидоров', 'Петров', 'Сазонов', 'Пирогов', 'Соболев', 'Титов', 'Максимов',
                      'Чукарин', 'Захаров']
    students = {}
    for i in range(n):
        gen = random.randint(0, 1)
        namenum = random.randint(0, 10)
        snamenum = random.randint(0, 10)
        zach_num = random.randint(0, n * 100)
        while zach_num in students.keys():
            zach_num = random.randint(0, n * 100)
        students[f'{zach_num} {student_names[gen][namenum]} {student_snames[snamenum]}{"а" if gen else ""}'] = random.randint(0, 400)
    return students


def save_to_sqlite(table):
    connector = sqlite3.connect('DataBase\PySQL1Labn.sql')
    cursor = connector.cursor()
    cursor.execute("""
                   CREATE TABLE IF NOT EXISTS students (
                       id INTEGER PRIMARY KEY AUTOINCREMENT, 
                       zachet_number STRING, 
                       name STRING, 
                       surname STRING,  
                       result STRING
                       )
                       """)
    cursor.execute("DELETE FROM students")
    cursor.executemany("INSERT INTO students (zachet_number, name, surname, result) VALUES (?, ?, ?, ?)", table)
    connector.commit()
    connector.close()


def get_from_sqlite():
    connector = sqlite3.connect('DataBase\PySQL1Labn.sql')
    cursor = connector.cursor()
    cursor.execute("SELECT zachet_number, name, surname, result FROM students")
    students = [[row[0], row[1], row[2], row[3]] for row in cursor.fetchall()]
    return students


def find_best_chet_num(students):
    lastres = -1
    lasti = -1
    for i, res in enumerate(students):
        if res[-1] % 2 == 0 and res[-1] > lastres:
            lastres = res[-1]
            lasti = i
    return lasti


def form_new(students, best_chet_num):
    answer = []
    for student in students:
        if student[-1] < best_chet_num:
            answer.append(student)

    return answer


def save_to_word(headers, tablew, text):
    path = 'DataBase\PyWord1Labn.docx'

    if os.path.exists(path):
        doc = Document(path)
    else:
        doc = Document()

    doc.add_heading(f'{text}:', level=1)
    table = doc.add_table(rows=len(tablew)+1, cols=len(headers))

    for index, header in enumerate(headers):
        table.cell(0, index).text = header

    for r, row in enumerate(tablew):
        for c, data in enumerate(row):
            table.cell(r + 1, c).text = data

    doc.add_paragraph()
    doc.save(path)


def save_to_excel(headers, tablew, text):
    path = 'DataBase\PyExcel1Labn.xlsx'
    if os.path.exists(path):
        wb = load_workbook(path)
    else:
        wb = Workbook()

    ws = wb.active
    ws.append([text])
    ws.append(headers)
    for row in tablew:
        ws.append(row)
    ws.append(['-----'])
    wb.save(path)


def save_to_pdf(headers, tablew, text, pdf):
    pdf.add_page()
    pdf.add_font('Arial', '', os.path.join('C:\\Windows\\Fonts', 'arial.ttf'), uni=True)
    pdf.set_font('Arial', size=12)
    pdf.cell(200, 10, text, ln=True, align='C')
    pdf.ln(5)
    pdf.cell(200, 10, '   '.join(headers), ln=True)

    for r, row in enumerate(tablew):
        pdf.cell(200, 10, '   '.join(row), ln=True)


def save_to_csv(headers, tablew, text):
    path = 'DataBase\PyCSV1Labn.csv'
    with open(path, mode='a', newline='', encoding='utf-16') as file:
        writer = csv.writer(file)
        writer.writerow([text])
        writer.writerow(headers)
        for row in tablew:
            writer.writerow(row)
        writer.writerow([])


def dict_tolist(dictionary):
    result = []
    for key, value in dictionary.items():
        result.append(key.split(' '))
        result[-1].append(str(value))
    return result


def write_anywhere(students, mcb, result):
    headers = ['Номер зачетной книжки', 'Фамилия', 'Имя', 'Сумма баллов']
    write_here = input('Сохранить в: \n'
                       'Word: w, Excel: e, SQLite: s, PDF: p, CSV: c\n'
                       'Пример ввода: wp (сохранение в Word, PDF)\n'
                       'Ваш ввод: ').lower()
    res = ''
    if 'w' in write_here:
        save_to_word(headers, students, 'Сгенерированный список студентов')
        save_to_word(headers, result, 'Результирующий список студентов')
        res += 'Word, '
    if 'e' in write_here:
        save_to_excel(headers, students, 'Сгенерированный список студентов')
        save_to_excel(headers, result, 'Результирующий список студентов')
        res += 'Excel, '
    if 'p' in write_here:
        path = 'DataBase\PyPDF1Labn.pdf'
        pdf = FPDF()
        save_to_pdf(headers, students, 'Сгенерированный список студентов', pdf)
        save_to_pdf(headers, result, 'Результирующий список студентов',  pdf)
        pdf.output(path)
        res += 'PDF, '
    if 'c' in write_here:
        save_to_csv(headers, students, 'Сгенерированный список студентов')
        save_to_csv(headers, result, 'Результирующий список студентов')
        res += 'CSV, '

    if not res:
        print('Запись не произведена')
    else:
        print('Записано в ' + res[:-2])


def lab_1n_start():
    print('Выполнение лабораторной работы номер 1')
    n = int(input('Введите количество студентов: '))
    genstudents = create_data(n)
    save_to_sqlite(dict_tolist(genstudents))
    students = get_from_sqlite()
    best_chet_num = find_best_chet_num(students)
    answer = form_new(students, best_chet_num)
    for i, row in enumerate(answer):
        for j, data in enumerate(row):
            if isinstance(data, int):
                answer[i][j] = str(data)
    for i, row in enumerate(students):
        for j, data in enumerate(row):
            if isinstance(data, int):
                students[i][j] = str(data)

    print('Список студентов:')
    for student in students:
        print(*student)
    print(f'Максимальный четный балл у студента номер {best_chet_num} ({students[best_chet_num]}):')
    print('Новый список студентов:')
    for student in answer:
        print(*student)

    write_anywhere(students, best_chet_num, answer)
