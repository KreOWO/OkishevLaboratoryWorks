import math as m
import random
from openpyxl import Workbook, load_workbook
from docx import Document
import sqlite3
import os


def vvod(text):
    """
    Запрашивает ввод определенного значения и конвертирует его в число
    :param text: Запрашиваемое значение
    :return: float от ввода
    """
    while True:
        try:
            return float(input(f'{text}: ').replace(',', '.'))
        except:
            print('Введите число!')


def vivod(var, res):
    """
    Выводит в консоль название переменной и её значение
    :param var: название переменной
    :param res: значение переменной
    """
    print(f'Переменная {var} = {res}')


def solve1(a, b, x):
    z = m.sqrt(a * x * m.sin(2 * x) + m.pow(m.e, (-2 * x)) * (x + b))
    return z


def solve2(zw, a, b, x):
    zw['w'] = pow(m.cos(pow(x, 3)), 2) - x / m.sqrt(pow(a, 2) + pow(b, 2))


def solve12(zw, a, b, x):
    zw['z'] = solve1(a, b, x)
    solve2(zw, a, b, x)


def gen_array(arr_len: int, minv: int, maxv: int):
    """
    Генерирует список случайных целых чисел
    :param arr_len: длина списка
    :param minv: минимальное значение
    :param maxv: максимальное значение
    :return: список случайных целых чисел
    """
    return [random.randint(minv, maxv) for _ in range(arr_len)]


def save_to_file(filename, text, value):
    """
    Сохраняет данные в файл.
    :param filename: имя файла
    :param text: подпись над данными
    :param value: данные
    """
    with open(filename, 'a', encoding='utf-8') as file:
        file.write(f"{text}:\n")
        if isinstance(value, list):
            file.write(", ".join(map(str, value)) + "\n\n")
        else:
            file.write(f"{value}\n\n")


def save_to_excel(filename, text, value):
    """
    Записывает данные в excel
    :param filename: путь к excel таблице
    :param text: текст над данными
    :param value: данные
    """
    if os.path.exists(filename):
        wb = load_workbook(filename)
    else:
        wb = Workbook()

    ws = wb.active
    ws.append([])

    ws.append([text])
    if isinstance(value, list):
        ws.append(value)
    else:
        ws.append([value])
    wb.save(filename)


def save_to_word(filename, text, value):
    """
    Создаёт word файл с таблицей
    :param filename: путь к word документу
    :param text: текст над данными
    :param value: данные
    """
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()

    doc.add_heading(f'{text}:', level=1)

    if isinstance(value, list):
        doc.add_paragraph(', '.join(map(str, value)))
    else:
        doc.add_paragraph(str(value))
    doc.save(filename)


def save_to_sqlite(filename, name, value):
    """
    Сохраняет данные в таблицу sql
    :param filename: имя файла
    :param name: название столбца
    :param value: данные
    """

    connector = sqlite3.connect(filename)
    connector.cursor().execute("CREATE TABLE IF NOT EXISTS results (id INTEGER PRIMARY KEY AUTOINCREMENT)")
    cursor = connector.cursor()

    cursor.execute(f"PRAGMA table_info(results)")
    columns = [row[1] for row in cursor.fetchall()]
    if name not in columns:
        cursor.execute(f"ALTER TABLE results ADD COLUMN {name} TEXT")

    cursor.execute("SELECT COUNT(*) FROM results")
    row_count = cursor.fetchone()[0]

    if isinstance(value, list):
        for i in range(min(row_count, len(value))):
            cursor.execute(f"UPDATE results SET {name} = {value[i]} WHERE rowid = {i + 1}")

        for i in range(row_count, len(value)):
            cursor.execute(f"INSERT INTO results ({name}) VALUES ({value[i]})")
    else:
        if row_count:
            cursor.execute(f"UPDATE results SET {name} = {value} WHERE rowid = 1")
        else:
            cursor.execute(f"INSERT INTO results ({name}) VALUES ({value})")

    connector.commit()
    connector.close()
